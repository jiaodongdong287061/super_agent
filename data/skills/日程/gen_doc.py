from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    run.element.rPr.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Microsoft YaHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()
    return table

# ============ 正文开始 ============

add_heading('SKILL.md 与系统提示词写作指南', level=0)

add_para('本文档详细介绍了如何编写 SKILL.md（技能说明书）和系统提示词（System Prompt），包括核心结构、写作原则、提炼规则和实战检查清单。')

# ============ 第一部分 ============

add_heading('一、SKILL.md 怎么写', level=1)

add_heading('1. 理解 SKILL.md 的本质', level=2)

add_para('SKILL.md 是一份给 AI Agent 读的技能说明书，它定义了：')
add_bullet('这个 Agent 能做什么（能力边界）')
add_bullet('什么时候触发（激活条件）')
add_bullet('怎么做（工作流程、模板、规范）')
add_bullet('什么不能做（注意事项、前置条件）')

add_heading('2. SKILL.md 的核心结构', level=2)

add_para('一个完整的 SKILL.md 应该包含以下模块，按顺序排列：')

add_code('---')
add_code('name: 技能名称')
add_code('description: 一句话描述')
add_code('---')
add_code('')
add_code('# 技能名称')
add_code('')
add_code('## 概述          <- 是什么、核心特色、前置条件')
add_code('## 触发条件      <- 什么情况下激活')
add_code('## 核心能力      <- 能做什么（分层描述）')
add_code('## 工作流程      <- 怎么做（步骤化）')
add_code('## 文件存储规范  <- 数据怎么存、怎么命名')
add_code('## 内置模板      <- 输出格式模板')
add_code('## 使用示例      <- 典型场景演示')
add_code('## 交互原则      <- 行为准则（优先级排序）')
add_code('## 注意事项      <- 禁止/限制/风险')
add_code('## 附录          <- 速查表、快速参考')

add_heading('3. 每个模块的写法要点', level=2)

add_heading('概述 — 最重要，Agent 最先读到的', level=3)

add_para('写法示例：', bold=True)

add_code('## 概述')
add_code('')
add_code('这是一个 XXX 技能，通过 XXX 方式，实现 XXX 功能。')
add_code('')
add_code('核心特色：')
add_code('- 特色1')
add_code('- 特色2')
add_code('')
add_code('前置条件（强制检查）：     <- 有前置条件一定要写在这里')
add_code('在执行 XXX 前必须 XXX...')

add_para('关键原则：不要把前置条件藏在文档中间，放在概述里，Agent 才能第一时间读到。', bold=True)

add_heading('触发条件 — 明确激活词', level=3)

add_code('## 触发条件')
add_code('')
add_code('当用户提到以下意图时触发：')
add_code('- "关键词1"、"关键词2"相关请求')
add_code('- "关键词3"')

add_para('关键原则：列出用户实际会说的词，不要写抽象描述。写"生成日报"而不是"执行日报告知性输出"。', bold=True)

add_heading('核心能力 — 用分层结构', level=3)

add_code('## 核心能力')
add_code('')
add_code('### 1. 状态感知          <- 第一层：大类')
add_code('- 文档交互：...          <- 第二层：具体方式')
add_code('- 问答交互：...')
add_code('')
add_code('### 2. 日程管理')
add_code('- 创建、修改、删除...')

add_para('关键原则：先大类再细分，每条用动词开头，说清楚做什么和怎么做。', bold=True)

add_heading('工作流程 — 用步骤或流程图', level=3)

add_para('两种写法：')

add_para('步骤式（适合线性流程）：', bold=True)
add_code('1. 收集任务：列出所有待办')
add_code('2. 评估优先级：按紧急程度排序')
add_code('3. 用户确认：获得确认后执行')

add_para('流程图式（适合有分支的流程）：', bold=True)
add_code('前置检查 -> 定时触发 -> 推送消息')
add_code('    |')
add_code('用户回复')
add_code('|- 继续 -> ...')
add_code('|- 静默 -> ...')

add_para('关键原则：流程中必须包含前置检查步骤和用户确认步骤，这是 Agent 最容易跳过的环节。', bold=True)

add_heading('文件存储规范 — 三要素', level=3)

add_code('## 文件存储规范')
add_code('')
add_code('目录结构：          <- 存在哪')
add_code('文件命名规则：      <- 怎么命名')
add_code('数据管理规则：      <- 覆盖还是新建、是否隔离')

add_heading('交互原则 — 按优先级排序', level=3)

add_code('## 交互原则')
add_code('')
add_code('1. **前置必检**：...     <- 最高优先级放第一')
add_code('2. **主动感知**：...')
add_code('3. **简洁确认**：...')

add_para('关键原则：第1条永远是最重要的、最容易违反的规则。', bold=True)

add_heading('注意事项 — 写"禁止"不写"建议"', level=3)

add_para('正确写法：', bold=True)
add_code('- 日程变更需获得用户明确确认        <- 明确')
add_code('- 敏感信息注意保护                  <- 明确')

add_para('错误写法：', bold=True)
add_code('- 建议在变更前确认                   <- 模糊，Agent 会忽略')
add_code('- 尽量注意信息安全                   <- 模糊，Agent 会忽略')

add_heading('4. SKILL.md 的核心写作原则', level=2)

add_table(
    ['原则', '说明', '反例'],
    [
        ['指令式而非描述式', '写"必须先确认"而非"建议确认"', '"建议先检查环境"'],
        ['流程化而非说明化', '写"第1步做什么->第2步做什么"', '"环境检查是重要的"'],
        ['穷举触发词', '列出用户会说的实际关键词', '"相关请求"'],
        ['写死分支路径', '每种用户回答都有明确的后续动作', '"根据情况处理"'],
        ['前置条件前置', '放在概述里，不要藏在文档中间', '把前置条件放在注意事项里'],
        ['模板内嵌', '把输出模板直接写在文档里', '"参考外部模板文件"'],
    ]
)

# ============ 第二部分 ============

add_heading('二、系统提示词怎么写', level=1)

add_heading('1. 理解系统提示词的本质', level=2)

add_para('系统提示词是 SKILL.md 的精简可执行版本。')
add_para('SKILL.md 是完整参考文档，系统提示词是 Agent 每次对话时加载的指令。')

add_table(
    ['', 'SKILL.md', '系统提示词'],
    [
        ['定位', '完整参考文档', '精简可执行指令'],
        ['受众', '人看的', 'Agent 运行时加载的'],
        ['内容', '完整、详细、含脚本和示例', '精简、只保留运行时需要的'],
        ['脚本/示例', '保留', '不保留'],
        ['配置详解', '完整步骤', '只保留关键 API 和参数'],
    ]
)

add_heading('2. 系统提示词的核心结构', level=2)

add_code('# 角色定义              <- 你是谁')
add_code('')
add_code('## 前置条件（最高优先级） <- 最重要，放最前面')
add_code('')
add_code('## 触发条件              <- 什么时候激活')
add_code('')
add_code('## 核心能力              <- 能做什么')
add_code('')
add_code('## 文件存储规范          <- 数据怎么存')
add_code('')
add_code('## 报告模板              <- 输出格式（精简版）')
add_code('')
add_code('## 工作流程              <- 怎么做（精简版）')
add_code('')
add_code('## 交互原则              <- 行为准则')
add_code('')
add_code('## 注意事项              <- 限制和风险')

add_heading('3. SKILL.md 到系统提示词的提炼规则', level=2)

add_table(
    ['SKILL.md 中的内容', '系统提示词中', '处理方式'],
    [
        ['概述', '保留', '精简为1-2句角色定义'],
        ['触发条件', '保留', '列出关键词即可'],
        ['核心能力', '保留', '每条精简到1行'],
        ['工作流程', '保留', '流程图改为箭头链式描述'],
        ['文件存储规范', '保留', '只保留目录结构和关键规则'],
        ['报告模板', '保留', '精简版，去掉可选字段'],
        ['交互原则', '保留', '原样保留，这是行为准则'],
        ['注意事项', '保留', '原样保留，这是红线'],
        ['工作流脚本', '不保留', '脚本是参考，Agent 运行时不需要'],
        ['定时任务配置示例', '合并', '只保留 rrule 速查表和调用方式'],
        ['使用示例', '不保留', '流程描述已覆盖'],
        ['配置步骤详解', '合并', '只保留关键 API 地址和请求示例'],
    ]
)

add_heading('4. 系统提示词的写作要点', level=2)

add_para('要点1：前置条件放最前面，加最高优先级标注', bold=True)

add_code('## 前置条件（强制检查，最高优先级）')

add_para('Agent 是从上往下读的，位置越靠前权重越高。')

add_para('要点2：把"文档说明"变成"可执行指令"', bold=True)

add_para('错误 — 说明式：')
add_code('必须安装 WorkBuddy，未安装时功能不可用')

add_para('正确 — 可执行式：')
add_code('检查方式：询问用户"你是否已配置 PushPlus？"')
add_code('- 已配置 -> 继续执行')
add_code('- 未配置 -> 告知"请先配置"，可继续但推送不可用')

add_para('要点3：每种用户回答都有明确分支', bold=True)

add_code('根据用户回答决定后续行为：')
add_code('- 已配置 -> 正常执行')
add_code('- 未配置 -> 告知需要配置，降级执行')
add_code('- 拒绝提供 -> 告知功能不可用')

add_para('不要留"根据情况处理"这种模糊分支。')

add_para('要点4：关键规则重复强调', bold=True)

add_para('重要的规则要在至少3个地方出现：')
add_bullet('概述/前置条件（首次强调）')
add_bullet('工作流程中（执行时强调）')
add_bullet('注意事项中（底线强调）')

add_para('例如"前置必检"就出现在了前置条件、工作流程、交互原则、注意事项四个地方。')

# ============ 第三部分 ============

add_heading('三、实战检查清单', level=1)

add_para('写完 SKILL.md 或系统提示词后，用以下清单自查：')

add_heading('SKILL.md 自查清单', level=2)

checklist_skill = [
    '概述中是否有前置条件？',
    '触发条件是否列出了用户实际会说的关键词？',
    '核心能力是否用动词开头？',
    '工作流程中是否包含前置检查和用户确认步骤？',
    '文件存储是否有目录结构 + 命名规则 + 管理规则？',
    '交互原则是否按优先级排序？第1条是否最重要？',
    '注意事项是否用"必须"而非"建议"？',
    '是否有分支场景（如未配置/已配置）的明确处理路径？',
]
for item in checklist_skill:
    add_bullet(item)

add_heading('系统提示词自查清单', level=2)

checklist_prompt = [
    '前置条件是否在最前面且标注"最高优先级"？',
    '检查流程是否是可执行的（问什么 -> 答什么 -> 做什么）？',
    '是否去掉了脚本和配置示例等参考性内容？',
    '关键规则是否在至少3个地方出现？',
    '是否每种用户回答都有明确的后续动作？',
    '模板是否只保留了必填字段？',
]
for item in checklist_prompt:
    add_bullet(item)

# ============ 总结 ============

add_heading('四、核心心法', level=1)

add_para('SKILL.md 是给人看的说明书，系统提示词是给 Agent 看的执行指令。写 SKILL.md 要完整详尽，写系统提示词要精简可执行。最关键的原则是：把"文档说明"变成"流程指令"，Agent 不会自己理解暗示，你必须把每一步都写死。', bold=True)

output_path = r'd:\workspace\jdd\创新项目组\IT运维数字员工\skills\日程\SKILL与系统提示词写作指南.docx'
doc.save(output_path)
print(f'Done: {output_path}')
