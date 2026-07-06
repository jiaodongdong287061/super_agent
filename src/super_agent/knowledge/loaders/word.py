import subprocess
import tempfile
from pathlib import Path

from langchain_core.documents import Document

from super_agent.knowledge.loaders.base import BaseLoader


class WordLoader(BaseLoader):
    def load(self, source: str) -> list[Document]:
        path = Path(source)
        if path.suffix.lower() == ".doc":
            return self._load_doc(path)
        return self._load_docx(source)

    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def _load_docx(self, source: str) -> list[Document]:
        import docx

        doc = docx.Document(source)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=full_text, metadata={"source": source})]

    def _load_doc(self, path: Path) -> list[Document]:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmpdir,
                    str(path.resolve()),
                ],
                check=True,
                capture_output=True,
                text=True,
                timeout=30,
            )
            converted = Path(tmpdir) / f"{path.stem}.docx"
            if not converted.exists():
                raise RuntimeError(
                    f"LibreOffice conversion failed for '{path}': "
                    f"expected '{converted}' not found"
                )
            return self._load_docx(str(converted))
